from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_UNDERLINE
from docx.shared import Inches, Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from schemas.employee import EmployeeRecord
from services.offer_letter.field_definitions import (
    BLANK,
    CLOSING_PARAGRAPH,
    FIELD_DEFINITIONS,
    OPENING_PARAGRAPH,
    appointment_ref,
    indian_long_date,
)
from io import BytesIO
from typing import Optional
from services.offer_letter.letterhead import get_footer_bytes, get_header_bytes

FONT = "Arial"
BODY_PT = Pt(10)
TITLE_PT = Pt(14)
HEADING_PT = Pt(11)


def _set_run_font(run, *, bold=False, size=BODY_PT, underline=False):
    run.bold = bold
    run.font.name = FONT
    run.font.size = size
    run._element.rPr.rFonts.set(qn("w:eastAsia"), FONT)
    if underline:
        run.font.underline = WD_UNDERLINE.SINGLE


def _add_horizontal_rule(paragraph, color="2E7D32"):
    p = paragraph._p
    p_pr = p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _add_header_footer_images(
    doc: Document,
    header_stream: Optional[BytesIO] = None,
    footer_stream: Optional[BytesIO] = None,
    company_name: Optional[str] = None
) -> None:
    section = doc.sections[0]
    content_width = section.page_width - section.left_margin - section.right_margin

    if header_stream is None:
        header_stream = get_header_bytes()

    if header_stream:
        header = section.header
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        run = header_para.add_run()
        run.add_picture(header_stream, width=content_width)
    else:
        hp = section.header.paragraphs[0] if section.header.paragraphs else section.header.add_paragraph()
        run = hp.add_run(company_name or "PaperlessBoss Private Limited")
        _set_run_font(run, bold=True)


def generate_appointment_docx(
    employee: EmployeeRecord,
    output_path: str | Path,
    header_bytes: Optional[BytesIO] = None,
    footer_bytes: Optional[BytesIO] = None,
    signature_image: Optional[str] = None,
    stamp_image: Optional[str] = None,
    company_name: Optional[str] = None
) -> Path:
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    today = date.today()
    ref_no = appointment_ref(employee.id, today.year)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.71)
    section.right_margin = Inches(0.71)

    _add_header_footer_images(doc, header_bytes, footer_bytes, company_name=company_name)

    def add_body_paragraph():
        return doc.add_paragraph()

    p = add_body_paragraph()
    r1 = p.add_run("Date: ")
    _set_run_font(r1, bold=True)
    r2 = p.add_run(indian_long_date(today))
    _set_run_font(r2)

    p = add_body_paragraph()
    r1 = p.add_run("Ref: ")
    _set_run_font(r1, bold=True)
    r2 = p.add_run(ref_no)
    _set_run_font(r2)

    p = add_body_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("APPOINTMENT LETTER")
    _set_run_font(run, bold=True, size=TITLE_PT, underline=True)

    hr = add_body_paragraph()
    _add_horizontal_rule(hr, color="2E7D32")

    company_name = employee.company_name or "PaperlessBoss Private Limited"
    p = add_body_paragraph()
    run = p.add_run("TERMS OF APPOINTMENT")
    _set_run_font(run, bold=True, size=HEADING_PT)

    hr = add_body_paragraph()
    _add_horizontal_rule(hr, color="646464")

    for field in FIELD_DEFINITIONS:
        value = field.getter(employee) or BLANK
        p = add_body_paragraph()
        r_roman = p.add_run(f"{field.roman}. ")
        _set_run_font(r_roman, bold=True)
        r_label = p.add_run(f"{field.label}: ")
        _set_run_font(r_label, bold=True)
        r_val = p.add_run(value)
        _set_run_font(r_val)

    hr = add_body_paragraph()
    _add_horizontal_rule(hr, color="646464")

    p = add_body_paragraph()
    run = p.add_run(CLOSING_PARAGRAPH)
    _set_run_font(run)

    p = add_body_paragraph()
    run = p.add_run(f"For {company_name}")
    _set_run_font(run, bold=True)

    if signature_image or stamp_image:
        import base64
        from reportlab.lib.units import Inches as RLInches # not needed, we can import from docx.shared
        from docx.shared import Inches as DocxInches
        p = add_body_paragraph()
        if signature_image:
            try:
                b64_sig = signature_image.split(";base64,")[1] if ";base64," in signature_image else signature_image
                sig_bytes = base64.b64decode(b64_sig)
                run = p.add_run()
                run.add_picture(BytesIO(sig_bytes), width=DocxInches(1.5))
            except Exception:
                pass
        if stamp_image:
            try:
                b64_stamp = stamp_image.split(";base64,")[1] if ";base64," in stamp_image else stamp_image
                stamp_bytes = base64.b64decode(b64_stamp)
                p.add_run("        ")
                run = p.add_run()
                run.add_picture(BytesIO(stamp_bytes), width=DocxInches(0.8))
            except Exception:
                pass
    else:
        doc.add_paragraph()

    p = add_body_paragraph()
    run = p.add_run("Authorised Signatory")
    _set_run_font(run, bold=True)

    p = add_body_paragraph()
    run = p.add_run("Name & Designation: ___________________________")
    _set_run_font(run)

    p = add_body_paragraph()
    run = p.add_run("Date: _____________________")
    _set_run_font(run)

    doc.add_paragraph()
    p = add_body_paragraph()
    run = p.add_run("ACKNOWLEDGEMENT BY EMPLOYEE")
    _set_run_font(run, bold=True, size=HEADING_PT)

    hr = add_body_paragraph()
    _add_horizontal_rule(hr, color="646464")

    p = add_body_paragraph()
    run = p.add_run(
        f"I, {employee.employee_name}, acknowledge receipt of this Appointment Letter "
        "and confirm my acceptance of all terms and conditions stated above."
    )
    _set_run_font(run)

    doc.add_paragraph()
    p = add_body_paragraph()
    run = p.add_run("Signature of Employee: ___________________________")
    _set_run_font(run)

    p = add_body_paragraph()
    run = p.add_run("Date: _____________________")
    _set_run_font(run)

    doc.core_properties.title = f"Appointment Letter – {employee.employee_name}"
    doc.core_properties.author = company_name
    doc.save(str(output_path))
    return output_path
